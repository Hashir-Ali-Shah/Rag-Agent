from pathlib import Path
from typing import List, Optional, Union, BinaryIO, TextIO
import logging
import io

# Optional imports with error handling
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DocumentReader:
    """
    Reads text from various document types: TXT, PDF, DOCX.
    Accepts both file paths (strings) and file objects from frontend.
    Returns a list of strings (paragraphs/pages).
    """
    
    def __init__(self):
        # Dynamic supported extensions based on available libraries
        self.supported_extensions = [".txt"]
        if PDF_AVAILABLE:
            self.supported_extensions.append(".pdf")
        if DOCX_AVAILABLE:
            self.supported_extensions.append(".docx")
    
    @property
    def SUPPORTED_EXTENSIONS(self) -> List[str]:
        return self.supported_extensions
    
    def read(self, file_input: Union[str, BinaryIO, TextIO], 
             filename: Optional[str] = None, 
             encoding: str = "utf-8") -> List[str]:
        """
        Read text from a document file or file object.
        
        Args:
            file_input: Either a file path (str) or file object from frontend
            filename: Original filename (required when using file objects to determine type)
            encoding: Text encoding (default: utf-8)
            
        Returns:
            List of text chunks (paragraphs for DOCX/TXT, pages for PDF)
            
        Raises:
            FileNotFoundError: If file path doesn't exist
            ValueError: If file type is unsupported or library unavailable
            Exception: For other reading errors
        """
        
        # Handle file path input
        if filename:
            return self._read_from_path(file_input, encoding)
        
        # Handle file object input
        else:
            if not filename:
                raise ValueError("filename parameter is required when using file objects")
            return self._read_from_file_object(file_input, filename, encoding)
    
    def _read_from_path(self, file_path: str, encoding: str) -> List[str]:
        """Read from file path (original functionality)."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
            
        ext = path.suffix.lower()
        
        if ext not in self.supported_extensions:
            self._check_library_availability(ext)
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.supported_extensions}")
        
        try:
            if ext == ".txt":
                return self._read_txt_from_path(path, encoding)
            elif ext == ".pdf":
                return self._read_pdf_from_path(path)
            elif ext == ".docx":
                return self._read_docx_from_path(path)
        except Exception as e:
            logging.error(f"Error reading {file_path}: {e}")
            raise Exception(f"Failed to read {file_path}: {e}")
        
        return []
    
    def _read_from_file_object(self, file_obj, filename: str, encoding: str) -> List[str]:
        """Read from file object."""
        ext = Path(filename).suffix.lower()
        
        if ext not in self.supported_extensions:
            self._check_library_availability(ext)
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.supported_extensions}")
        
        try:
            if ext == ".txt":
                return self._read_txt_from_object(file_obj, encoding)
            elif ext == ".pdf":
                return self._read_pdf_from_object(file_obj)
            elif ext == ".docx":
                return self._read_docx_from_object(file_obj)
        except Exception as e:
            logging.error(f"Error reading {filename}: {e}")
            raise Exception(f"Failed to read {filename}: {e}")
        
        return []
    
    def _check_library_availability(self, ext: str):
        """Check if required library is available for file type."""
        if not PDF_AVAILABLE and ext == ".pdf":
            raise ValueError("PDF support requires PyPDF2 library")
        if not DOCX_AVAILABLE and ext == ".docx":
            raise ValueError("DOCX support requires python-docx library")
    
    # Path-based reading methods (original)
    def _read_txt_from_path(self, path: Path, encoding: str) -> List[str]:
        """Read text file from path and split by double newlines."""
        try:
            with open(path, "r", encoding=encoding) as f:
                content = f.read()
                texts = [p.strip() for p in content.split("\n\n") if p.strip()]
                return texts if texts else [content.strip()] if content.strip() else []
        except UnicodeDecodeError:
            # Fallback to different encodings
            for fallback_encoding in ['latin-1', 'cp1252', 'utf-16']:
                try:
                    with open(path, "r", encoding=fallback_encoding) as f:
                        content = f.read()
                        texts = [p.strip() for p in content.split("\n\n") if p.strip()]
                        return texts if texts else [content.strip()] if content.strip() else []
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode text file with any common encoding")
    
    def _read_pdf_from_path(self, path: Path) -> List[str]:
        """Read PDF file from path and extract text from each page."""
        if not PDF_AVAILABLE:
            raise ValueError("PDF support requires PyPDF2 library")
            
        reader = PdfReader(str(path))
        return self._extract_pdf_text(reader)
    
    def _read_docx_from_path(self, path: Path) -> List[str]:
        """Read DOCX file from path and extract text from each paragraph."""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX support requires python-docx library")
            
        doc = docx.Document(str(path))
        return self._extract_docx_text(doc)
    
    # File object-based reading methods (new)
    def _read_txt_from_object(self, file_obj, encoding: str) -> List[str]:
        """Read text from file object."""
        try:
            # Reset file pointer to beginning
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)
            
            # Handle different file object types
            if hasattr(file_obj, 'read'):
                if isinstance(file_obj, io.TextIOBase):
                    # Text file object
                    content = file_obj.read()
                else:
                    # Binary file object
                    content = file_obj.read()
                    if isinstance(content, bytes):
                        try:
                            content = content.decode(encoding)
                        except UnicodeDecodeError:
                            # Try fallback encodings
                            for fallback_encoding in ['latin-1', 'cp1252', 'utf-16']:
                                try:
                                    content = content.decode(fallback_encoding)
                                    break
                                except UnicodeDecodeError:
                                    continue
                            else:
                                raise ValueError("Could not decode text file with any common encoding")
                
                texts = [p.strip() for p in content.split("\n\n") if p.strip()]
                return texts if texts else [content.strip()] if content.strip() else []
            else:
                raise ValueError("File object does not support read() method")
                
        except Exception as e:
            raise Exception(f"Failed to read text from file object: {e}")
    
    def _read_pdf_from_object(self, file_obj) -> List[str]:
        """Read PDF from file object and extract text from each page."""
        if not PDF_AVAILABLE:
            raise ValueError("PDF support requires PyPDF2 library")
        
        try:
            # Reset file pointer to beginning
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)
            
            # PyPDF2 can handle file-like objects directly
            reader = PdfReader(file_obj)
            return self._extract_pdf_text(reader)
            
        except Exception as e:
            raise Exception(f"Failed to read PDF from file object: {e}")
    
    def _read_docx_from_object(self, file_obj) -> List[str]:
        """Read DOCX from file object and extract text from each paragraph."""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX support requires python-docx library")
        
        try:
            # Reset file pointer to beginning
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)
            
            # python-docx can handle file-like objects directly
            doc = docx.Document(file_obj)
            return self._extract_docx_text(doc)
            
        except Exception as e:
            raise Exception(f"Failed to read DOCX from file object: {e}")
    
    # Helper methods for text extraction
    def _extract_pdf_text(self, reader) -> List[str]:
        """Extract text from PDF reader object."""
        texts = []
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    texts.append(page_text.strip())
            except Exception as e:
                logging.warning(f"Could not extract text from page {page_num + 1}: {e}")
                continue
        return texts
    
    def _extract_docx_text(self, doc) -> List[str]:
        """Extract text from DOCX document object."""
        texts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                texts.append(para.text.strip())
        return texts

# Example usage and testing
if __name__ == "__main__":
    from pathlib import Path

    # Initialize reader
    reader = DocumentReader()
    print(f"Supported extensions: {reader.SUPPORTED_EXTENSIONS}")

    # Your file path
    file_path = r"D:\resume\current.pdf"  # replace with your file

    # Open file as a file object
    ext = Path(file_path).suffix.lower()

    mode = "rb" if ext in [".pdf", ".docx"] else "r"  # PDF/DOCX need binary mode

    with open(file_path, mode) as file_obj:
        texts = reader.read(file_obj, filename=Path(file_path).name)

    # Print results
    for i, text in enumerate(texts):
        print(f"Chunk {i+1} ({len(text)} chars):\n{text}\n")
